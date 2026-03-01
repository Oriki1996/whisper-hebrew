"""core/docx_exporter.py — Academic DOCX export with full Hebrew RTL support."""
from typing import Optional


def _fmt_time(sec: float) -> str:
    h = int(sec) // 3600
    m = int(sec) // 60 % 60
    s = int(sec) % 60
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _set_rtl(paragraph) -> None:
    """Make a python-docx paragraph right-to-left."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _add_heading(doc, text: str, level: int = 1, color_hex: str = "2b1f14") -> None:
    from docx.shared import Pt, RGBColor
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    _set_rtl(p)


def _add_paragraph(doc, text: str, bold: bool = False,
                   italic: bool = False, size_pt: int = 11,
                   color_hex: str = "2b1f14") -> None:
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color_hex)
    _set_rtl(p)


def generate_docx(
    lecture: dict,
    segments: list[dict],
    insights: Optional[dict] = None,
    entities: Optional[dict] = None,
) -> bytes:
    """
    Generate an academic-style DOCX with proper Hebrew RTL support.

    Args:
        lecture:  Lecture row dict.
        segments: List of segment dicts.
        insights: Optional insights dict.
        entities: Optional NER entities dict.

    Returns:
        DOCX file as raw bytes.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io
    except ImportError:
        raise ImportError("pip install python-docx")

    import io as _io

    doc = Document()

    # ── Page margins (RTL document) ─────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = int(21.0 * 914400 / 25.4)   # A4 in EMUs
    section.page_height = int(29.7 * 914400 / 25.4)
    section.left_margin  = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin   = Cm(2.5)
    section.bottom_margin= Cm(2.5)

    # Set document direction to RTL
    settings = doc.settings.element
    doc_settings = OxmlElement("w:bidi")
    settings.append(doc_settings)

    # ── Title ───────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    title_run = title_p.add_run(lecture.get("filename", "הרצאה"))
    title_run.font.size  = Pt(20)
    title_run.bold       = True
    title_run.font.color.rgb = RGBColor.from_string("c8813a")
    title_p.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl(title_p)

    # ── Metadata ────────────────────────────────────────────────────────────
    meta_pairs = [
        ("קורס",  lecture.get("course_name")),
        ("מרצה",  lecture.get("lecturer")),
        ("תאריך", lecture.get("date")),
        ("משך",   _fmt_time(lecture.get("duration", 0)) if lecture.get("duration") else None),
    ]
    for label, val in meta_pairs:
        if val:
            p = doc.add_paragraph()
            r_label = p.add_run(f"{label}: ")
            r_label.bold = True
            r_label.font.size = Pt(10)
            r_label.font.color.rgb = RGBColor.from_string("6b5740")
            r_val = p.add_run(val)
            r_val.font.size = Pt(10)
            r_val.font.color.rgb = RGBColor.from_string("6b5740")
            _set_rtl(p)

    doc.add_paragraph()  # spacer

    # ── Summary ─────────────────────────────────────────────────────────────
    if insights and insights.get("summary"):
        _add_heading(doc, "סיכום", level=1, color_hex="2b1f14")
        _add_paragraph(doc, insights["summary"])
        doc.add_paragraph()

    # ── Key Terms ───────────────────────────────────────────────────────────
    if insights and insights.get("key_terms"):
        _add_heading(doc, "מושגי מפתח", level=1, color_hex="2b1f14")
        for t in insights["key_terms"]:
            p = doc.add_paragraph(style="List Bullet")
            r_term = p.add_run(t.get("term", ""))
            r_term.bold = True
            r_term.font.size = Pt(10)
            p.add_run(f" — {t.get('definition', '')}").font.size = Pt(10)
            _set_rtl(p)
        doc.add_paragraph()

    # ── Transcript ──────────────────────────────────────────────────────────
    if segments:
        _add_heading(doc, "תמלול", level=1, color_hex="2b1f14")
        current_speaker = None
        for seg in segments:
            spk = seg.get("speaker_id", "")
            if spk and spk != current_speaker:
                _add_heading(doc, spk.replace("SPEAKER_", "דובר "), level=2,
                             color_hex="5b8dd9")
                current_speaker = spk

            ts   = _fmt_time(seg.get("start_time", 0))
            text = seg.get("text", "").strip()
            p    = doc.add_paragraph()

            r_ts = p.add_run(f"[{ts}]  ")
            r_ts.font.size  = Pt(8)
            r_ts.font.color.rgb = RGBColor.from_string("6b5740")

            r_text = p.add_run(text)
            r_text.font.size = Pt(10)
            _set_rtl(p)

        doc.add_paragraph()

    # ── References ──────────────────────────────────────────────────────────
    def _has_refs():
        if entities:
            return any(entities.get(k) for k in ("authors", "books", "laws", "cases"))
        return bool(insights and insights.get("citations"))

    if _has_refs():
        _add_heading(doc, "רשימת מקורות", level=1, color_hex="2b1f14")
        n = 1

        if entities:
            for author in entities.get("authors", []):
                field = f" ({author['field']})" if author.get("field") else ""
                tss   = ", ".join(_fmt_time(t) for t in author.get("timestamps", [])[:3])
                ts_s  = f" [{tss}]" if tss else ""
                _add_paragraph(doc, f"{n}. {author.get('name','')}{field}{ts_s}")
                n += 1

            for book in entities.get("books", []):
                a_s  = f" — {book['author']}" if book.get("author") else ""
                y_s  = f" ({book['year']})" if book.get("year") else ""
                tss  = ", ".join(_fmt_time(t) for t in book.get("timestamps", [])[:3])
                ts_s = f" [{tss}]" if tss else ""
                _add_paragraph(doc, f"{n}. {book.get('title','')}{a_s}{y_s}{ts_s}",
                               italic=True)
                n += 1

            for law in entities.get("laws", []):
                y_s  = f" ({law['year']})" if law.get("year") else ""
                tss  = ", ".join(_fmt_time(t) for t in law.get("timestamps", [])[:3])
                ts_s = f" [{tss}]" if tss else ""
                _add_paragraph(doc, f"{n}. {law.get('name','')}{y_s}{ts_s}")
                n += 1

            for case in entities.get("cases", []):
                c_s  = f" ({case['court']})" if case.get("court") else ""
                y_s  = f" {case['year']}" if case.get("year") else ""
                tss  = ", ".join(_fmt_time(t) for t in case.get("timestamps", [])[:3])
                ts_s = f" [{tss}]" if tss else ""
                _add_paragraph(doc, f"{n}. {case.get('name','')}{c_s}{y_s}{ts_s}",
                               italic=True)
                n += 1

        elif insights and insights.get("citations"):
            for c in insights["citations"]:
                ref = c.get("author", "")
                if c.get("title"): ref += f", {c['title']}"
                if c.get("year"):  ref += f" ({c['year']})"
                _add_paragraph(doc, f"{n}. {ref}")
                n += 1

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
